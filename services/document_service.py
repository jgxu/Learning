import os
import json
from sqlalchemy.orm import Session
from fastapi import UploadFile
from models import Document, User
from config import settings
import services.ai_service as ai_service
from datetime import datetime

# 确保上传目录存在
os.makedirs(settings.UPLOAD_DIR, exist_ok=True)


def upload_document(db: Session, openid: str, file: UploadFile = None, link: str = None) -> Document:
    """上传文件或链接并返回解析文本"""
    user = db.query(User).filter(User.openid == openid).first()
    if not user:
        raise ValueError("用户不存在")
    
    document = Document(user_id=user.id)
    
    if file:
        # 处理文件上传
        file_ext = file.filename.split('.')[-1].lower()
        file_path = os.path.join(settings.UPLOAD_DIR, f"{datetime.utcnow().timestamp()}_{file.filename}")
        
        # 保存文件
        with open(file_path, "wb") as buffer:
            buffer.write(file.file.read())
        
        document.title = file.filename
        document.file_type = get_file_type(file_ext)
        document.file_path = file_path
        
        # 解析文件内容
        content = parse_file(file_path, document.file_type)
        document.content = content
        document.is_parsed = True
        
    elif link:
        # 处理链接
        document.title = link.split('/')[-1] if '/' in link else link
        document.file_type = "link"
        document.source_link = link
        
        # 解析链接内容
        content = parse_link(link)
        document.content = content
        document.is_parsed = True
    
    db.add(document)
    db.commit()
    db.refresh(document)
    
    return document


def get_document_history(db: Session, openid: str, page: int = 1, page_size: int = 10) -> list[Document]:
    """分页获取上传历史"""
    user = db.query(User).filter(User.openid == openid).first()
    if not user:
        raise ValueError("用户不存在")
    
    offset = (page - 1) * page_size
    documents = db.query(Document).filter(Document.user_id == user.id)
    documents = documents.order_by(Document.created_at.desc()).offset(offset).limit(page_size).all()
    
    return documents


def get_file_type(extension: str) -> str:
    """根据文件扩展名获取文件类型"""
    if extension in ["doc", "docx"]:
        return "word"
    elif extension == "pdf":
        return "pdf"
    elif extension == "txt":
        return "txt"
    elif extension in ["jpg", "jpeg", "png", "gif"]:
        return "image"
    else:
        return "other"


def parse_file(file_path: str, file_type: str) -> str:
    """解析文件内容"""
    try:
        if file_type == "word":
            return parse_word_file(file_path)
        elif file_type == "pdf":
            return parse_pdf_file(file_path)
        elif file_type == "txt":
            return parse_txt_file(file_path)
        elif file_type == "image":
            return parse_image_file(file_path)
        else:
            return "不支持的文件类型"
    except Exception as e:
        return f"解析失败: {str(e)}"


def parse_word_file(file_path: str) -> str:
    """解析Word文件"""
    try:
        from docx import Document
        doc = Document(file_path)
        content = "\n".join([para.text for para in doc.paragraphs])
        return content
    except ImportError:
        return "请安装python-docx库"


def parse_pdf_file(file_path: str) -> str:
    """解析PDF文件"""
    try:
        import PyPDF2
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            content = "\n".join([page.extract_text() for page in reader.pages])
        return content
    except ImportError:
        return "请安装PyPDF2库"


def parse_txt_file(file_path: str) -> str:
    """解析TXT文件"""
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()


def parse_image_file(file_path: str) -> str:
    """解析图片文件（使用Gemini API）"""
    try:
        return ai_service.analyze_image(file_path)
    except Exception as e:
        return f"图片解析失败: {str(e)}"


def parse_link(link: str) -> str:
    """解析链接内容"""
    try:
        import requests
        from bs4 import BeautifulSoup
        
        response = requests.get(link)
        soup = BeautifulSoup(response.text, "html.parser")
        
        # 提取主要内容
        paragraphs = soup.find_all("p")
        content = "\n".join([p.get_text() for p in paragraphs])
        
        return content[:10000]  # 限制内容长度
    except Exception as e:
        return f"链接解析失败: {str(e)}"